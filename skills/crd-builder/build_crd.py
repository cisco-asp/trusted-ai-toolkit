#!/usr/bin/env python3
"""Generic CRD builder — produces CRD_<Customer>_<Project>_<YYYYMMDD>.docx from template.docx.

Usage (called by the agent after populating the DATA dict below):

    python build_crd.py

The agent must edit the DATA dict at the bottom of this file with customer-specific
content before running the script. Every value is a plain string unless noted.
"""

from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path

# ---------------------------------------------------------------------------
# Resolve python-docx: prefer system install, fall back to vendored _pydeps/
# ---------------------------------------------------------------------------
_skill_dir = Path(__file__).resolve().parent
_pydeps_candidates = [
    Path.cwd() / "_pydeps",
    _skill_dir / "_pydeps",
]
for _p in _pydeps_candidates:
    if _p.exists():
        sys.path.insert(0, str(_p))
        break

from docx import Document  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.text.paragraph import Paragraph  # noqa: E402

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SKILL_DIR = Path(__file__).resolve().parent
TEMPLATE = SKILL_DIR / "template.docx"

# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------


def _out_path(customer: str, project: str, out_dir: Path) -> Path:
    tag = f"{customer}_{project}".replace(" ", "_")
    return out_dir / f"CRD_{tag}_{datetime.now().strftime('%Y%m%d')}.docx"


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)


def insert_paragraph_before(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def remove_red_runs(doc: Document) -> None:
    """Erase guideline text rendered in red (FF0000) by the template."""
    for p in doc.paragraphs:
        for run in p.runs:
            rpr = run._r.find(qn("w:rPr"))
            if rpr is not None:
                color_el = rpr.find(qn("w:color"))
                if (
                    color_el is not None
                    and color_el.get(qn("w:val"), "").upper() == "FF0000"
                ):
                    run.text = ""


def strip_instruction_pages(doc: Document) -> None:
    """Delete paragraphs that belong to template instruction pages.

    The template marks removal pages with the string below.  We also delete
    everything that precedes the formal title page paragraph
    "Customer Requirements Document".
    """
    MARKER = "<< REMOVE THIS PAGE FROM FINAL DOCUMENT >>"
    TITLE = "Customer Requirements Document"

    paragraphs = list(doc.paragraphs)
    # Walk forward; once we see the title stop deleting.
    deleting = True
    for p in paragraphs:
        if MARKER in p.text:
            deleting = True
        if p.text.strip() == TITLE:
            deleting = False
        if deleting and p.text.strip() != TITLE:
            delete_paragraph(p)


def _replace_prompt(doc: Document, prompt: str, replacement: str) -> None:
    """Find a paragraph whose text matches *prompt* exactly and replace its content."""
    for p in doc.paragraphs:
        if p.text.strip() == prompt:
            p.clear()
            p.add_run(replacement)
            return


def _insert_extra_sections(doc: Document, sections: list[tuple[str, str]]) -> None:
    """Insert (heading, body) pairs as Heading 3 + normal paragraphs before §5.

    *sections* is a list of (heading_text, body_text) in the desired reading order.
    Pass an empty list to skip.
    """
    if not sections:
        return

    anchor5 = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("5.") and "Design Considerations" in p.text:
            anchor5 = p
            break
    if anchor5 is None:
        raise RuntimeError("Could not find §5 Design Considerations heading")

    for heading, body in reversed(sections):
        ph = insert_paragraph_before(anchor5, heading)
        ph.style = "Heading 3"
        body_el = OxmlElement("w:p")
        ph._p.addnext(body_el)
        pb = Paragraph(body_el, ph._parent)
        pb.add_run(body)
        anchor5 = ph


def _add_gap_table(
    doc: Document,
    heading_text: str,
    rows: list[tuple[str, ...]],
    col_headers: list[str],
) -> None:
    """Insert a requirements-gap table (or any multi-column table) before §5.

    *rows* is a list of tuples; each tuple length must equal len(col_headers).
    The table is inserted immediately after the paragraph whose text matches
    *heading_text* (which must already exist in the document, e.g. inserted via
    _insert_extra_sections).
    """
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == heading_text:
            anchor = p
            break
    if anchor is None:
        raise RuntimeError(f"Heading not found for gap table: {heading_text!r}")

    tbl = doc.add_table(rows=1 + len(rows), cols=len(col_headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(col_headers):
        hdr[i].text = h
    for ri, row_data in enumerate(rows, start=1):
        for ci, val in enumerate(row_data):
            tbl.rows[ri].cells[ci].text = str(val)
    # Detach from end of document and place after the heading paragraph
    tbl._tbl.getparent().remove(tbl._tbl)
    anchor._p.addnext(tbl._tbl)


def _duplicate_feature_block(doc: Document) -> None:
    """Append a copy of the Feature Details + Use Case tables before §5.

    Call once per additional feature (beyond the first).  The copies are
    appended in document order; fill them after calling this function by
    accessing the newly appended tables at the end of doc.tables.
    """
    anchor5 = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("5.") and "Design Considerations" in p.text:
            anchor5 = p
            break
    if anchor5 is None:
        raise RuntimeError("Could not find §5 heading for feature block duplication")

    # tables[5], [6], [7] = feature detail, use-case 1, use-case 2
    for tbl in (doc.tables[5], doc.tables[6], doc.tables[7]):
        clone = deepcopy(tbl._tbl)
        anchor5._p.addprevious(clone)


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------


def build(data: dict) -> Path:
    """Produce the CRD .docx file from *data* and return the output path."""

    out_dir = Path(data.get("out_dir", Path.cwd()))
    out_dir.mkdir(parents=True, exist_ok=True)
    out = _out_path(data["customer"], data["project"], out_dir)

    shutil.copy(TEMPLATE, out)
    doc = Document(str(out))

    strip_instruction_pages(doc)
    remove_red_runs(doc)

    today = datetime.now().strftime("%m/%d/%Y")

    # ------------------------------------------------------------------
    # Cover table
    # ------------------------------------------------------------------
    t0 = doc.tables[0]
    t0.rows[0].cells[1].text = data["customer"]
    t0.rows[1].cells[1].text = data["product_platform"]
    t0.rows[2].cells[1].text = data["technology_solution"]
    t0.rows[3].cells[1].text = data["author"]
    t0.rows[4].cells[1].text = data.get("date", today)

    # ------------------------------------------------------------------
    # Reviewers table
    # ------------------------------------------------------------------
    t1 = doc.tables[1]
    reviewers = data.get("reviewers", [])  # list of (department, name_title)
    for i, (dept, name) in enumerate(reviewers):
        row_idx = i + 1
        if row_idx >= len(t1.rows):
            # Add a new row by copying the last data row's XML
            last_tr = t1.rows[-1]._tr
            new_tr = deepcopy(last_tr)
            last_tr.addnext(new_tr)
        t1.rows[row_idx].cells[0].text = dept
        t1.rows[row_idx].cells[1].text = name

    # ------------------------------------------------------------------
    # Modification history table
    # ------------------------------------------------------------------
    t2 = doc.tables[2]
    t2.rows[1].cells[0].text = "1.0"
    t2.rows[1].cells[1].text = data.get("date", today)
    t2.rows[1].cells[2].text = data["author"]
    t2.rows[1].cells[3].text = data.get("mod_history_comments", "Initial release.")

    # ------------------------------------------------------------------
    # Executive overview prompts
    # ------------------------------------------------------------------
    overview = data.get("overview", {})
    _replace_prompt(
        doc,
        "Provide a description that covers what the purpose of this request?",
        overview.get("purpose", ""),
    )
    _replace_prompt(
        doc,
        "What are the problems it will solve?",
        overview.get("problems", ""),
    )
    _replace_prompt(
        doc,
        "Provide a description of the problems that this request is attempting to solve?",
        overview.get("technical_detail", ""),
    )

    # ------------------------------------------------------------------
    # Goals and Success Criteria table
    # ------------------------------------------------------------------
    t3 = doc.tables[3]
    goals = data.get("goals", [])  # list of (description, target_date, measurement)
    for i, (desc, target, measure) in enumerate(goals[:3]):
        row_idx = i + 1
        t3.rows[row_idx].cells[1].text = desc
        t3.rows[row_idx].cells[2].text = target
        t3.rows[row_idx].cells[3].text = measure

    # ------------------------------------------------------------------
    # Business case table
    # ------------------------------------------------------------------
    t4 = doc.tables[4]
    bc = data.get("business_case", {})
    t4.rows[0].cells[2].text = bc.get("customer", data["customer"])
    t4.rows[1].cells[2].text = bc.get("region", "")
    t4.rows[2].cells[2].text = bc.get("segment", "")
    t4.rows[3].cells[2].text = bc.get("contact", "")
    t4.rows[4].cells[2].text = bc.get("advanced_services", "")

    # ------------------------------------------------------------------
    # Features (one or more)
    # ------------------------------------------------------------------
    features = data.get("features", [])
    # The template ships with one feature block (tables 5, 6, 7).
    # For each additional feature beyond the first, duplicate the block.
    for _ in range(len(features) - 1):
        _duplicate_feature_block(doc)

    FEATURE_ROWS = [
        "Feature Name",
        "Aha Idea #",
        "Customer Feature ID",
        "Contract Section",
        "Feature Description",
        "Platforms",
        "Critical Feature?",
        "Scalability and Performance Expectations Details",
        "Golden Architecture Model",
        "Acceptance Criteria",
        "Solution Validation",
        "HW Requirements and Dependencies",
        "Backward Compatibility",
        "Deployment Plan",
    ]

    for feat_idx, feat in enumerate(features):
        tbl_offset = feat_idx * 3  # each feature occupies 3 tables: detail + 2 use-case
        t_detail = doc.tables[5 + tbl_offset]
        t_uc1 = doc.tables[6 + tbl_offset]
        t_uc2 = doc.tables[7 + tbl_offset]

        # Header row (merged cell spanning both columns)
        header = feat.get("title", f"Feature {feat_idx + 1} – Details")
        t_detail.rows[0].cells[0].text = header
        t_detail.rows[0].cells[1].text = header

        detail_values = [
            feat.get("name", ""),
            feat.get("aha_id", ""),
            feat.get("customer_feature_id", ""),
            feat.get("contract_section", ""),
            feat.get("description", ""),
            feat.get("platforms", ""),
            feat.get("critical", ""),
            feat.get("scalability", ""),
            feat.get("golden_architecture", ""),
            feat.get("acceptance_criteria", ""),
            feat.get("solution_validation", ""),
            feat.get("hw_requirements", ""),
            feat.get("backward_compatibility", ""),
            feat.get("deployment_plan", ""),
        ]
        for ri, val in enumerate(detail_values, start=1):
            t_detail.rows[ri].cells[1].text = val

        use_cases = feat.get("use_cases", ["", ""])
        t_uc1.rows[1].cells[0].text = use_cases[0] if len(use_cases) > 0 else ""
        t_uc2.rows[1].cells[0].text = use_cases[1] if len(use_cases) > 1 else ""

    # ------------------------------------------------------------------
    # Optional extra narrative sections (inserted before §5)
    # ------------------------------------------------------------------
    _insert_extra_sections(doc, data.get("extra_sections", []))

    # ------------------------------------------------------------------
    # Optional requirements gap table (inserted after a named heading)
    # ------------------------------------------------------------------
    gap_table = data.get("gap_table")
    if gap_table:
        _add_gap_table(
            doc,
            heading_text=gap_table["heading"],
            rows=gap_table["rows"],
            col_headers=gap_table["col_headers"],
        )

    # ------------------------------------------------------------------
    # Design Considerations Checklist
    # ------------------------------------------------------------------
    chk = None
    for tbl in doc.tables:
        if tbl.rows[0].cells[0].text.startswith("High Availability"):
            chk = tbl
            break
    if chk:
        checklist = data.get("checklist", {})
        # checklist keys: row index (int) -> {"yn": "Y", "eft": ""}
        for row_idx, vals in checklist.items():
            row = chk.rows[int(row_idx)]
            if "yn" in vals:
                row.cells[1].text = vals["yn"]
            if "eft" in vals:
                row.cells[3].text = vals["eft"]

    # ------------------------------------------------------------------
    # Appendix (§6) — source list paragraph
    # ------------------------------------------------------------------
    sources = data.get("appendix_sources", [])
    if sources:
        for p in doc.paragraphs:
            if p.text.strip().startswith("6.") and "Appendix" in p.text:
                insert_paragraph_before(p, "Sources: " + "; ".join(sources))
                break

    # ------------------------------------------------------------------
    # Optional appendix spreadsheet table (before "End of Document")
    # ------------------------------------------------------------------
    appendix_table_rows = data.get("appendix_table_rows")  # list[list[str]]
    appendix_table_heading = data.get("appendix_table_heading", "Appendix Table")
    if appendix_table_rows:
        p_end = None
        for p in doc.paragraphs:
            if p.text.strip() == "End of Document":
                p_end = p
                break
        if p_end is None:
            raise RuntimeError('"End of Document" paragraph not found')
        ncols = max(len(r) for r in appendix_table_rows)
        p_intro = insert_paragraph_before(p_end, "")
        p_h = insert_paragraph_before(p_intro, appendix_table_heading)
        p_h.style = "Heading 3"
        tbl = doc.add_table(rows=len(appendix_table_rows), cols=ncols)
        tbl.style = "Table Grid"
        tbl._tbl.getparent().remove(tbl._tbl)
        p_intro._p.addnext(tbl._tbl)
        for ri, row in enumerate(appendix_table_rows):
            for ci, val in enumerate(row):
                if ci < len(tbl.rows[ri].cells):
                    tbl.rows[ri].cells[ci].text = str(val)

    doc.save(str(out))
    print(f"Wrote {out}")
    return out


# ---------------------------------------------------------------------------
# DATA — agent populates this dict before running the script
# ---------------------------------------------------------------------------
# fmt: off
DATA = {
    # ---- Required --------------------------------------------------------
    "customer":            "ACME Corp",           # Customer name
    "project":             "SRv6_Core",           # Short project tag (used in filename)
    "author":              "userid@cisco.com",     # Cisco UserID
    "product_platform":    "ASR 9000",             # Product / Platform
    "technology_solution": "SRv6 uSID Core",      # Technology / Solution

    # ---- Optional --------------------------------------------------------
    "date":                "",          # Leave blank to use today's date (MM/DD/YYYY)
    "out_dir":             ".",         # Output directory path

    "mod_history_comments": (
        "Initial CRD."
    ),

    "overview": {
        "purpose": (
            "Describe the purpose of this request."
        ),
        "problems": (
            "Describe the problems this feature will solve."
        ),
        "technical_detail": (
            "Provide technical detail on the problems being addressed."
        ),
    },

    # List of (goal_description, target_date, measurement) — up to 3 rows
    "goals": [
        ("Goal 1 description.", "Q4 CY2026", "Acceptance criterion or metric."),
        ("Goal 2 description.", "Q1 CY2027", "Acceptance criterion or metric."),
        ("Goal 3 description.", "TBD",        "Acceptance criterion or metric."),
    ],

    "business_case": {
        "customer":          "ACME Corp",
        "region":            "North America",
        "segment":           "Service Provider",
        "contact":           "Jane Smith (SE); John Doe (AM)",
        "advanced_services": "Yes — CX engagement planned.",
    },

    # List of reviewer tuples: (department, name/title)
    "reviewers": [
        ("Account / Sales",     "John Doe (Account Manager)"),
        ("Systems Engineering", "Jane Smith (SE)"),
        ("Architecture",        "Alice Brown"),
        ("CX / PM",             "Bob Lee (SSEM)"),
    ],

    # One dict per feature.  For a single feature use a list with one item.
    "features": [
        {
            "title":                  "Feature 1 – Details",
            "name":                   "Feature name",
            "aha_id":                 "IOSXR-XXXXX",
            "customer_feature_id":    "Customer-facing ID or N/A",
            "contract_section":       "TBD",
            "description":            "What the feature does and why the customer needs it.",
            "platforms":              "Cisco 8000; ASR 9000",
            "critical":               "Yes",
            "scalability":            "Scale targets and performance expectations.",
            "golden_architecture":    "Reference architecture or topology description.",
            "acceptance_criteria":    "How success will be measured.",
            "solution_validation":    "Lab test, CPOC, or field trial reference.",
            "hw_requirements":        "Hardware models, optics, or dependencies.",
            "backward_compatibility": "Cross-release behavior and migration notes.",
            "deployment_plan":        "Rollout plan and maintenance window notes.",
            # Two use-case narrative strings
            "use_cases": [
                "Use case 1 narrative.",
                "Use case 2 narrative.",
            ],
        },
        # Add more feature dicts here for additional features, e.g.:
        # {
        #     "title": "Feature 2 – Details",
        #     ...
        # },
    ],

    # Optional: extra Heading 3 + body sections inserted before §5.
    # List of (heading, body) in desired reading order.  Leave empty to skip.
    "extra_sections": [
        # ("Customer requirements summary", "Narrative text..."),
        # ("Lab / validation results",      "Narrative text..."),
        # ("Platform caveats",              "Narrative text..."),
    ],

    # Optional: requirements gap table inserted after a named heading above.
    # Set to None to skip.
    "gap_table": None,
    # Example:
    # "gap_table": {
    #     "heading":     "Customer requirements summary",   # must match an extra_sections heading
    #     "col_headers": ["Requirement", "Target", "Aha ID", "Gap / Status"],
    #     "rows": [
    #         ("Requirement text", "Target or N/A", "IOSXR-XXXXX", "Gap description"),
    #     ],
    # },

    # Design considerations checklist.
    # Keys are row indices (0-based) in the checklist table.
    # "yn"  -> column 1 (Y / N / N/A)
    # "eft" -> column 3 (EFT date or Y / N / N/A)
    "checklist": {
        0: {"yn": "Y", "eft": "Y"},
        1: {"yn": "Y", "eft": "Y"},
        2: {"yn": "Y", "eft": ""},
    },

    # List of source file descriptions for the §6 Appendix paragraph.
    "appendix_sources": [
        # "Customer requirements document v1.0 (PDF)",
        # "Lab validation report (PDF)",
        # "Feature requirements matrix (XLSX)",
    ],

    # Optional appendix table rows (list of lists of strings).
    # Set to None to skip.
    "appendix_table_rows":    None,
    "appendix_table_heading": "Appendix Table",
}
# fmt: on

if __name__ == "__main__":
    build(DATA)
